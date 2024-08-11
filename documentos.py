from docx import Document
from docx.shared import Cm
import os
import subprocess
from time import sleep

cores = {
    'azul': '\033[94m',
    'vermelho': '\033[91m',
    'magenta': '\033[95m',
    'resetar': '\033[0m'
}

def gerar_word(cotacao, site, data, imagem):
    
    print(f'{cores['azul']}-> Criando o Arquivo no Formato Word!{cores['resetar']}')

    user_profile = os.getenv('USERPROFILE')
    desktop_path = os.path.join(user_profile, 'Desktop')
    arquivo = os.path.join(desktop_path, 'Cotacao.docx')

    if not os.path.exists(arquivo):
        documento = Document()
        documento.save(arquivo)
    
    documento = Document(arquivo)
    documento.add_heading(f'Cotação do Dolar R${cotacao} - ({data})', 0)

    paragrafo_um = documento.add_paragraph(f'O dólar foi cotado no valor de ')
    paragrafo_um.add_run(f'R$ {cotacao}').bold = True
    paragrafo_um.add_run(f', na data {data}.')

    documento.add_paragraph(f'O site utilizado para a cotação foi {site}')
    documento.add_picture(imagem, width=Cm(16.00))
    documento.add_paragraph(f'Cotação feita por ').add_run('Adilson Gustavo\n').bold
    
    for _ in range(4):
        documento.add_paragraph()

    documento.save(arquivo)

    print(f'{cores['azul']}---> Arquivo Criado Com Sucesso na Área de Trabalho!{cores['resetar']}')
    print(f'{cores['azul']}---> Nome do Arquivo: Cotacao.docx{cores['resetar']}')
    sleep(1)


def converte_pdf():

    print(f'{cores['vermelho']}-> Transformando o Arquivo Word em PDF!{cores['resetar']}')
        
    user_profile = os.getenv('USERPROFILE')
    desktop_path = os.path.join(user_profile, 'Desktop')
    docx_path = os.path.join(desktop_path, 'Cotacao.docx')
    soffice_path = 'C:\\Program Files\\LibreOffice\\program\\soffice.exe'

    command = [
        soffice_path,
        '--headless', 
        '--convert-to', 'pdf',
        '--outdir', desktop_path,
        docx_path
    ]

    try:
        result = subprocess.run(command, capture_output=True, text=True, check=True)
        print(f'{cores['vermelho']}---> Arquivo Transformado com Sucesso na Área de Trabalho!{cores['resetar']}')
        print(f'{cores['vermelho']}---> Nome do Arquivo: Cotacao.pdf{cores['resetar']}')
        sleep(1)
    except subprocess.CalledProcessError as e:
        print(f"{cores['magenta']}---> Não foi possível Transformar o Arquivo Word em PDF{cores['resetar']}")